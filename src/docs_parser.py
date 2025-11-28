import os
import uuid
import re
import shutil
from docx import Document
from docx.oxml import CT_Picture
from typing import Tuple, List, Dict
from src.config import IMAGES_DIR

def read_txt_md(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return text

def extract_images_from_docx(file_path: str, doc_id: str) -> Tuple[str, List[str]]:
    """
    Извлечение изображений из .docx файла

    Args:
        file_path: Путь к .docx файлу
        doc_id: ID документа для именования изображений

    Returns:
        Tuple (текст с плейсхолдерами, список путей к изображениям)
    """
    # Создаем директорию для изображений, если не существует
    os.makedirs(IMAGES_DIR, exist_ok=True)

    doc = Document(file_path)
    image_paths = []
    text_parts = []
    image_counter = 0

    for paragraph in doc.paragraphs:
        # Проверяем наличие изображений в параграфе
        for run in paragraph.runs:
            # Проверяем inline shapes (встроенные изображения)
            if 'graphic' in run._element.xml:
                image_counter += 1

                # Генерируем имя файла для изображения
                image_filename = f"{doc_id}_img_{image_counter}.png"
                image_path = os.path.join(IMAGES_DIR, image_filename)

                # Извлекаем изображение
                try:
                    # Получаем объект изображения из XML
                    blip = run._element.xpath('.//a:blip')
                    if blip:
                        embed = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed:
                            image = doc.part.related_parts[embed]
                            with open(image_path, 'wb') as f:
                                f.write(image.blob)

                            # Добавляем плейсхолдер
                            relative_path = os.path.relpath(image_path, start=os.path.dirname(os.path.dirname(__file__)))
                            image_paths.append(relative_path)
                            text_parts.append(f"[[image: {relative_path}]]")
                except Exception as e:
                    print(f"⚠️  Ошибка извлечения изображения: {e}")

        # Добавляем текст параграфа, если есть
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Объединяем текст
    text_with_placeholders = "\n".join(text_parts)

    return text_with_placeholders, image_paths

def extract_images_from_markdown(file_path: str, doc_id: str) -> Tuple[str, List[str]]:
    """
    Извлечение изображений из .md файла

    Args:
        file_path: Путь к .md файлу
        doc_id: ID документа для именования изображений

    Returns:
        Tuple (текст с плейсхолдерами, список путей к изображениям)
    """
    # Создаем директорию для изображений, если не существует
    os.makedirs(IMAGES_DIR, exist_ok=True)

    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    image_paths = []
    image_counter = 0

    # Регулярное выражение для поиска markdown изображений: ![alt](path)
    # Также поддерживает ![alt](path "title")
    image_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'

    def replace_image(match):
        nonlocal image_counter
        image_source = match.group(2).split()[0]  # Убираем title если есть

        # Если путь относительный, копируем изображение
        if not image_source.startswith(('http://', 'https://', 'data:')):
            source_dir = os.path.dirname(file_path)
            source_image_path = os.path.join(source_dir, image_source)

            if os.path.exists(source_image_path):
                image_counter += 1

                # Определяем расширение
                _, ext = os.path.splitext(source_image_path)
                if not ext:
                    ext = '.png'

                # Генерируем новое имя
                image_filename = f"{doc_id}_img_{image_counter}{ext}"
                dest_image_path = os.path.join(IMAGES_DIR, image_filename)

                # Копируем изображение
                try:
                    shutil.copy2(source_image_path, dest_image_path)
                    relative_path = os.path.relpath(dest_image_path, start=os.path.dirname(os.path.dirname(__file__)))
                    image_paths.append(relative_path)
                    return f"[[image: {relative_path}]]"
                except Exception as e:
                    print(f"⚠️  Ошибка копирования изображения {source_image_path}: {e}")
                    return match.group(0)  # Возвращаем исходный текст при ошибке
            else:
                print(f"⚠️  Изображение не найдено: {source_image_path}")
                return match.group(0)
        else:
            # Для URL оставляем как есть (можно добавить скачивание в будущем)
            return match.group(0)

    # Заменяем все изображения
    text_with_placeholders = re.sub(image_pattern, replace_image, text)

    return text_with_placeholders, image_paths

def extract_text(file_path: str) -> str:
    """
    Извлечение текста из файла

    Args:
        file_path: Путь к файлу

    Returns:
        Извлеченный текст
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt" or ext == ".md":
        return read_txt_md(file_path)
    elif ext == ".docx":
        return read_docx(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")

def extract_text_with_filename(file_path: str) -> Tuple[str, str]:
    """
    Извлечение текста из файла с названием файла

    Args:
        file_path: Путь к файлу

    Returns:
        Tuple (текст, название файла без расширения)
    """
    text = extract_text(file_path)
    filename = os.path.basename(file_path)
    # Убираем расширение из названия
    filename_without_ext = os.path.splitext(filename)[0]
    return text, filename_without_ext

def prepare_text_for_chunking(text: str, filename_without_ext: str) -> str:
    """
    Подготовка текста для разбиения на чанки с добавлением названия документа

    Args:
        text: Исходный текст документа
        filename_without_ext: Название файла без расширения

    Returns:
        Текст с добавленным названием документа
    """
    # Добавляем название документа в начало
    header = f"Документ: {filename_without_ext}\n\n"
    return header + text


# === Новые функции для работы с двумя типами документов ===

def parse_single_instruction(
    file_path: str,
    tags: List[str] = None,
    author: str = "Admin"
) -> List[Dict]:
    """
    Обработка файла с одной инструкцией
    Название файла = название инструкции

    Args:
        file_path: Путь к файлу
        tags: Список тегов
        author: Автор документа

    Returns:
        Список из одной инструкции с метаданными
    """
    filename_without_ext = os.path.splitext(os.path.basename(file_path))[0]
    file_format = os.path.splitext(file_path)[1][1:]  # без точки
    doc_id = str(uuid.uuid4())

    # Извлекаем текст и изображения в зависимости от формата
    images = []
    if file_format == 'docx':
        text, images = extract_images_from_docx(file_path, doc_id)
    elif file_format == 'md':
        text, images = extract_images_from_markdown(file_path, doc_id)
    else:
        text = extract_text(file_path)

    instruction_data = {
        'id': str(uuid.uuid4()),
        'doc_id': doc_id,
        'title': filename_without_ext,
        'file_path': file_path,
        'file_format': file_format,
        'source_type': 'single_file',
        'separator_index': None,
        'text': text,
        'tags': tags or [],
        'author': author,
        'images': images
    }

    return [instruction_data]


def parse_multi_instructions(
    file_path: str,
    tags: List[str] = None,
    author: str = "Admin"
) -> List[Dict]:
    """
    Обработка файла с множеством инструкций
    Разделитель: ---
    Первая строка после --- = название инструкции

    Args:
        file_path: Путь к файлу
        tags: Список тегов (применяются ко всем инструкциям)
        author: Автор документа

    Returns:
        Список инструкций с метаданными
    """
    file_format = os.path.splitext(file_path)[1][1:]
    doc_id = str(uuid.uuid4())  # Общий doc_id для всех инструкций

    # Извлекаем текст с обработкой изображений
    all_images = []
    if file_format == 'docx':
        text, all_images = extract_images_from_docx(file_path, doc_id)
    elif file_format == 'md':
        text, all_images = extract_images_from_markdown(file_path, doc_id)
    else:
        text = extract_text(file_path)

    # Разделение по ---
    sections = text.split('---')
    instructions = []

    for idx, section in enumerate(sections):
        section = section.strip()
        if not section:
            continue

        # Первая строка = название инструкции
        lines = section.split('\n', 1)
        title = lines[0].strip()
        content = lines[1].strip() if len(lines) > 1 else ""

        if not title:
            # Если нет названия, пропускаем
            continue

        # Находим изображения в этой секции
        section_images = [img for img in all_images if img in content]

        instruction_data = {
            'id': str(uuid.uuid4()),
            'doc_id': doc_id,
            'title': title,
            'file_path': file_path,
            'file_format': file_format,
            'source_type': 'multi_instruction',
            'separator_index': idx,
            'text': content,
            'tags': tags or [],
            'author': author,
            'images': section_images
        }

        instructions.append(instruction_data)

    return instructions


def parse_document(
    file_path: str,
    source_type: str,
    tags: List[str] = None,
    author: str = "Admin"
) -> List[Dict]:
    """
    Универсальная функция для парсинга документа

    Args:
        file_path: Путь к файлу
        source_type: Тип источника ('single_file' или 'multi_instruction')
        tags: Список тегов
        author: Автор документа

    Returns:
        Список инструкций с метаданными
    """
    if source_type == 'single_file':
        return parse_single_instruction(file_path, tags, author)
    elif source_type == 'multi_instruction':
        return parse_multi_instructions(file_path, tags, author)
    else:
        raise ValueError(f"Неподдерживаемый тип источника: {source_type}")
